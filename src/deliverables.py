from __future__ import annotations

import os
from pathlib import Path

import matplotlib.pyplot as plt
import numpy as np
import torch
from sklearn.model_selection import KFold, train_test_split


PROJECT_ROOT = Path(__file__).resolve().parents[1]
os.chdir(PROJECT_ROOT)

# Import after setting working directory so feedforward.py can load the npz reliably.
from feedforward import (  # noqa: E402
    FeedforwardNN,
    VariableDepthNN,
    X,
    y,
    fit_linear_regression_normal_equation,
    make_fold_dataloaders,
    train_model,
)


def _ensure_2d_col(v: np.ndarray) -> np.ndarray:
    v = np.asarray(v)
    if v.ndim == 1:
        return v.reshape(-1, 1)
    return v


def _rmse(a: np.ndarray, b: np.ndarray) -> float:
    return float(np.sqrt(np.mean((a - b) ** 2)))


def run_experiments(
    depths: list[list[int]] | None = None,
    n_splits: int = 5,
    epochs: int = 120,
    lr: float = 1e-3,
    batch_size: int = 32,
    random_state: int = 42,
) -> dict:
    if depths is None:
        depths = [[32], [32, 16], [32, 16, 8], [64, 32, 16, 8]]

    X_train, X_test, y_train, y_test = train_test_split(
        X, y, test_size=0.2, random_state=random_state
    )
    y_test_col = _ensure_2d_col(y_test)

    kf = KFold(n_splits=n_splits, shuffle=True, random_state=random_state)

    linear_rmses: list[float] = []
    nn_rmses: list[float] = []
    nn_histories: list[dict[str, list[float]]] = []
    nn_test_preds_all: list[np.ndarray] = []

    depth_histories: dict[str, list[dict[str, list[float]]]] = {
        str(d): [] for d in depths
    }
    depth_rmses: dict[str, list[float]] = {str(d): [] for d in depths}

    for tr_idx, va_idx in kf.split(X_train):
        # ----- Linear baseline (fit on normalized fold-train, eval on test in original units)
        X_tr, y_tr = X_train[tr_idx], _ensure_2d_col(y_train[tr_idx])
        X_mean = X_tr.mean(axis=0)
        X_std = np.where(X_tr.std(axis=0) == 0, 1.0, X_tr.std(axis=0))
        y_mean = y_tr.mean(axis=0)
        y_std = np.where(y_tr.std(axis=0) == 0, 1.0, y_tr.std(axis=0))

        X_tr_norm = (X_tr - X_mean) / X_std
        X_test_norm = (X_test - X_mean) / X_std
        y_tr_norm = (y_tr - y_mean) / y_std

        linear_rmse = fit_linear_regression_normal_equation(
            X_train_norm=X_tr_norm,
            y_train_norm=y_tr_norm,
            X_test_norm=X_test_norm,
            y_test_orig=y_test_col,
            y_mean=y_mean,
            y_std=y_std,
        )
        linear_rmses.append(linear_rmse)

        # ----- NN baseline depth [32,16,8]
        train_loader, val_loader, stats = make_fold_dataloaders(
            X_train, y_train, tr_idx, va_idx, batch_size=batch_size
        )

        model = FeedforwardNN(input_dim=X_train.shape[1], output_dim=1)
        history = train_model(
            model=model,
            train_loader=train_loader,
            val_loader=val_loader,
            epochs=epochs,
            lr=lr,
        )
        nn_histories.append(history)

        X_test_norm_t = torch.tensor(
            ((X_test - stats["X_mean"]) / stats["X_std"]), dtype=torch.float32
        )
        model.eval()
        with torch.no_grad():
            pred_norm = model(X_test_norm_t).cpu().numpy()

        pred = pred_norm * np.asarray(stats["y_std"]).reshape(1, -1) + np.asarray(
            stats["y_mean"]
        ).reshape(1, -1)
        nn_test_preds_all.append(pred)
        nn_rmses.append(_rmse(pred, y_test_col))

        # ----- Depth experiment
        for depth in depths:
            d_key = str(depth)
            d_model = VariableDepthNN(
                input_dim=X_train.shape[1], hidden_dims=depth, output_dim=1
            )
            d_history = train_model(
                model=d_model,
                train_loader=train_loader,
                val_loader=val_loader,
                epochs=epochs,
                lr=lr,
            )
            depth_histories[d_key].append(d_history)

            d_model.eval()
            with torch.no_grad():
                d_pred_norm = d_model(X_test_norm_t).cpu().numpy()
            d_pred = d_pred_norm * np.asarray(stats["y_std"]).reshape(
                1, -1
            ) + np.asarray(stats["y_mean"]).reshape(1, -1)
            depth_rmses[d_key].append(_rmse(d_pred, y_test_col))

    nn_pred_mean = np.mean(np.stack(nn_test_preds_all, axis=0), axis=0)

    return {
        "X_test": X_test,
        "y_test": y_test_col,
        "linear_rmses": linear_rmses,
        "nn_rmses": nn_rmses,
        "nn_histories": nn_histories,
        "nn_pred_mean": nn_pred_mean,
        "depth_histories": depth_histories,
        "depth_rmses": depth_rmses,
        "depths": depths,
        "epochs": epochs,
    }


def make_plots(results: dict, out_dir: Path) -> dict[str, Path]:
    out_dir.mkdir(parents=True, exist_ok=True)

    plot_paths: dict[str, Path] = {}

    # 1) CV loss curves (NN baseline)
    epochs = len(results["nn_histories"][0]["train_loss"])
    x = np.arange(1, epochs + 1)
    train_mat = np.array([h["train_loss"] for h in results["nn_histories"]])
    val_mat = np.array([h["val_loss"] for h in results["nn_histories"]])

    plt.figure(figsize=(8, 5))
    for row in train_mat:
        plt.plot(x, row, color="tab:blue", alpha=0.2)
    for row in val_mat:
        plt.plot(x, row, color="tab:orange", alpha=0.2)
    plt.plot(
        x, train_mat.mean(axis=0), color="tab:blue", linewidth=2.5, label="Train (mean)"
    )
    plt.plot(
        x, val_mat.mean(axis=0), color="tab:orange", linewidth=2.5, label="Val (mean)"
    )
    plt.xlabel("Epoch")
    plt.ylabel("MSE Loss")
    plt.title("Cross-Validation Loss Curves (NN 32-16-8)")
    plt.legend()
    plt.tight_layout()
    p1 = out_dir / "cv_loss_curves.png"
    plt.savefig(p1, dpi=160)
    plt.close()
    plot_paths["cv_loss"] = p1

    # 2) Predicted vs actual (NN mean prediction)
    y_true = results["y_test"].reshape(-1)
    y_pred = results["nn_pred_mean"].reshape(-1)
    min_v = min(y_true.min(), y_pred.min())
    max_v = max(y_true.max(), y_pred.max())

    plt.figure(figsize=(6.5, 6.0))
    plt.scatter(y_true, y_pred, alpha=0.7)
    plt.plot([min_v, max_v], [min_v, max_v], "r--", linewidth=2)
    plt.xlabel("Actual Picking Time")
    plt.ylabel("Predicted Picking Time")
    plt.title("Predicted vs Actual (NN)")
    plt.tight_layout()
    p2 = out_dir / "predicted_vs_actual.png"
    plt.savefig(p2, dpi=160)
    plt.close()
    plot_paths["scatter"] = p2

    # 3) Depth experiment curves (mean val loss per epoch)
    plt.figure(figsize=(8, 5))
    for d in results["depths"]:
        key = str(d)
        val_curves = np.array([h["val_loss"] for h in results["depth_histories"][key]])
        plt.plot(x, val_curves.mean(axis=0), linewidth=2.0, label=key)
    plt.xlabel("Epoch")
    plt.ylabel("Validation MSE")
    plt.title("Depth Experiment: Mean CV Validation Curves")
    plt.legend(title="Hidden layers")
    plt.tight_layout()
    p3 = out_dir / "depth_experiment_curves.png"
    plt.savefig(p3, dpi=160)
    plt.close()
    plot_paths["depth_curves"] = p3

    return plot_paths


def build_docx(results: dict, plot_paths: dict[str, Path], out_doc: Path) -> None:
    try:
        from docx import Document
        from docx.shared import Inches
    except ImportError as exc:
        raise ImportError(
            "python-docx is required to create the report. "
            "Install it in this environment with:\n"
            "  /Users/mattcooper/ME467/ME467_HW67/.venv/bin/python -m ensurepip --upgrade\n"
            "  /Users/mattcooper/ME467/ME467_HW67/.venv/bin/python -m pip install python-docx"
        ) from exc

    linear_mean = float(np.mean(results["linear_rmses"]))
    linear_std = float(np.std(results["linear_rmses"]))
    nn_mean = float(np.mean(results["nn_rmses"]))
    nn_std = float(np.std(results["nn_rmses"]))

    best_depth = min(
        results["depth_rmses"], key=lambda k: np.mean(results["depth_rmses"][k])
    )
    best_depth_rmse = float(np.mean(results["depth_rmses"][best_depth]))

    doc = Document()
    doc.add_heading("Exercise 1 Report", level=1)
    doc.add_paragraph(
        "This report summarizes cross-validation training behavior, baseline comparisons, "
        "depth experiments, and final test performance for the warehouse picking-time task."
    )

    doc.add_heading("1. Confirmation of implementation", level=2)
    doc.add_paragraph(
        "The PyTorch pipeline includes: cross-validation training, final test evaluation, "
        "linear normal-equation baseline, and a depth experiment with variable hidden layers."
    )

    doc.add_heading("2. CV loss curves", level=2)
    doc.add_picture(str(plot_paths["cv_loss"]), width=Inches(6.3))

    doc.add_heading("3. Predicted vs actual", level=2)
    doc.add_picture(str(plot_paths["scatter"]), width=Inches(5.8))

    doc.add_heading("4. Linear vs NN RMSE comparison", level=2)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Light Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Model"
    hdr[1].text = "Mean RMSE"
    hdr[2].text = "Std RMSE"

    row = table.add_row().cells
    row[0].text = "Linear baseline"
    row[1].text = f"{linear_mean:.4f}"
    row[2].text = f"{linear_std:.4f}"

    row = table.add_row().cells
    row[0].text = "NN (32-16-8)"
    row[1].text = f"{nn_mean:.4f}"
    row[2].text = f"{nn_std:.4f}"

    doc.add_heading("5. Depth experiment", level=2)
    doc.add_picture(str(plot_paths["depth_curves"]), width=Inches(6.3))

    depth_table = doc.add_table(rows=1, cols=3)
    depth_table.style = "Light Grid"
    dh = depth_table.rows[0].cells
    dh[0].text = "Hidden layers"
    dh[1].text = "Mean RMSE"
    dh[2].text = "Std RMSE"

    for depth_key, vals in results["depth_rmses"].items():
        r = depth_table.add_row().cells
        r[0].text = depth_key
        r[1].text = f"{np.mean(vals):.4f}"
        r[2].text = f"{np.std(vals):.4f}"

    doc.add_heading("6. Discussion", level=2)
    improvement = 100.0 * (linear_mean - nn_mean) / linear_mean
    doc.add_paragraph(
        "The neural network captures nonlinear interactions in the data that are difficult "
        "for linear regression to represent. "
        f"Compared with linear baseline, the 32-16-8 NN changed RMSE by {improvement:.2f}% "
        "on average. The depth experiment shows how capacity affects generalization; very "
        "shallow models can underfit while overly deep models may not improve validation loss. "
        f"The best tested depth was {best_depth} with mean RMSE {best_depth_rmse:.4f}."
    )

    out_doc.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_doc)


def main() -> None:
    print("Running experiments for deliverables...")
    results = run_experiments()

    artifacts_dir = PROJECT_ROOT / "artifacts"
    plot_paths = make_plots(results, artifacts_dir)

    report_path = PROJECT_ROOT / "exercise_1_report.docx"
    build_docx(results, plot_paths, report_path)

    print("\nDone.")
    print(f"Report: {report_path}")
    print(f"Plots dir: {artifacts_dir}")
    print(
        "Confirmation: includes CV, final test evaluation, linear baseline comparison, and depth experiment."
    )


if __name__ == "__main__":
    main()
